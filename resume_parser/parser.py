import pdfplumber
import docx


def extract_text_from_resume(file_path):

    """
    Extract text from PDF or DOCX resume
    """

    text = ""

    if file_path.endswith(".pdf"):
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""

    elif file_path.endswith(".docx"):
        document = docx.Document(file_path)
        for para in document.paragraphs:
            text += para.text

    return text

import docx


def check_font_consistency(file_path):

    fonts = []

    document = docx.Document(file_path)

    for para in document.paragraphs:
        for run in para.runs:
            if run.font.name:
                fonts.append(run.font.name)

    unique_fonts = set(fonts)

    if len(unique_fonts) <= 2:
        return 5

    return 2